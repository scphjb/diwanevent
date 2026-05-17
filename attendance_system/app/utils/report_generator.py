import os
from docx import Document
from io import BytesIO
from datetime import datetime

def generate_official_report(event_name, stats, participants):
    """
    توليد المحضر الرسمي للفعالية باستخدام قالب Word.
    """
    template_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "Eventool_Report_AR.docx")
    
    if not os.path.exists(template_path):
        # Fallback if template missing
        doc = Document()
        doc.add_heading(f"تقرير فعالية: {event_name}", 0)
    else:
        doc = Document(template_path)

    # استبدال الكلمات المفتاحية في المستند
    # ملاحظة: هذا يتطلب وجود {{EVENT_NAME}} و {{TOTAL}} إلخ في القالب
    replacements = {
        "{{EVENT_NAME}}": event_name,
        "{{TOTAL_INVITED}}": str(stats['total_invited']),
        "{{CHECKED_IN}}": str(stats['checked_in']),
        "{{ATTENDANCE_RATE}}": f"{stats['attendance_rate']}%",
        "{{DATE}}": datetime.now().strftime("%Y-%m-%d")
    }

    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    # إضافة جدول المشاركين
    doc.add_page_break()
    doc.add_heading("قائمة الحضور التفصيلية", level=1)
    
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'الرقم'
    hdr_cells[1].text = 'الاسم الكامل'
    hdr_cells[2].text = 'الجهة'
    hdr_cells[3].text = 'وقت الحضور'

    for i, p in enumerate(participants):
        if p.payment_status == 'paid':
            row_cells = table.add_row().cells
            row_cells[0].text = str(i + 1)
            row_cells[1].text = p.full_name
            row_cells[2].text = p.organization
            row_cells[3].text = p.attendance_records[0].check_in_time.strftime("%H:%M") if p.attendance_records else "---"

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
